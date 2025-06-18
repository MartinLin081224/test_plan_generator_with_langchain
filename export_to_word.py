from docx import Document
import markdown
from bs4 import BeautifulSoup
import os


def convert_markdown_to_word(md_path, word_path):
    """將 Markdown 檔案轉換為 Word 檔案"""
    if not os.path.exists(md_path):
        raise FileNotFoundError(f"❌ 找不到 Markdown 檔案：{md_path}")

    with open(md_path, "r", encoding="utf-8") as f:
        md_text = f.read()

    # 將 Markdown 轉為 HTML
    html = markdown.markdown(md_text, extensions=["tables"])
    soup = BeautifulSoup(html, "html.parser")

    doc = Document()

    for elem in soup.recursiveChildGenerator():
        if elem.name:
            if elem.name == "h1":
                doc.add_heading(elem.get_text(), level=1)
            elif elem.name == "h2":
                doc.add_heading(elem.get_text(), level=2)
            elif elem.name == "h3":
                doc.add_heading(elem.get_text(), level=3)
            elif elem.name == "p":
                doc.add_paragraph(elem.get_text())
            elif elem.name == "ul":
                continue  # 已透過 li 加入
            elif elem.name == "li":
                doc.add_paragraph("• " + elem.get_text(), style="List Bullet")
            elif elem.name == "table":
                rows = elem.find_all("tr")
                if not rows:
                    continue
                num_cols = len(rows[0].find_all(["td", "th"]))
                table = doc.add_table(rows=len(rows), cols=num_cols)
                table.style = "Table Grid"
                for i, row in enumerate(rows):
                    cells = row.find_all(["td", "th"])
                    for j, cell in enumerate(cells):
                        text = cell.get_text()
                        table.cell(i, j).text = text

    doc.save(word_path)
    print(f"✅ Word 已儲存至：{word_path}")


if __name__ == "__main__":
    import sys
    if len(sys.argv) != 3:
        print("用法：python export_to_word.py <input.md> <output.docx>")
    else:
        convert_markdown_to_word(sys.argv[1], sys.argv[2])
